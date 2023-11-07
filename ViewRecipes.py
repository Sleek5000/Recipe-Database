import tkinter as tk
from tkinter import *
from tkinter import Image
from tkinter.ttk import Combobox
import base64
import io
import PIL.Image
import mysql.connector
import win32api
import win32print
from PIL import ImageTk, Image
from docx import Document
from docx.shared import Inches
#import ViewImage
image_list = []


class ViewRecipesUI:

    def __init__(self):
        self.view_recipe_window = tk.Tk()
        self.view_recipe_window.title('Select Recipe')
        self.view_recipe_window.geometry('1920x1080')
        self.line1_frame = tk.Frame(self.view_recipe_window)

        # first row
        self.recipe_label = tk.Label(self.line1_frame, text="Search By Ingredient:")
        self.recipe_label.grid(row=0, column=0)
        self.search_by_ingredient_input = tk.Entry(self.line1_frame, width=50)
        self.search_by_ingredient_input.grid(row=0, column=1)
        self.search_button = tk.Button(self.line1_frame, text="Search", command=self.getRecipes)
        self.search_button.grid(row=0, column=2)

        # second row
        self.recipe_list = tk.StringVar()

        self.recipe_combo_box = Combobox(self.line1_frame, width=50, textvariable=self.recipe_list)
        self.recipe_combo_box['values'] = ['Select from list']
        self.recipe_combo_box.current(0)
        self.recipe_combo_box.grid(row=1, column=1)
        self.recipe_combo_box_label = tk.Label(self.line1_frame, text="List of Recipes")
        self.recipe_combo_box_label.grid(row=1, column=0)
        self.show_recipe_button = tk.Button(self.line1_frame,
                                            text="Show Recipe",
                                            command=lambda: [self.viewRecipe(), self.viewImage()])
        self.show_recipe_button.grid(row=1, column=2)

        # third row
        self.ingredients = tk.StringVar()
        self.ingredients_label = tk.Label(self.line1_frame, textvariable=self.ingredients)
        self.ingredients.set('Ingredients:')
        self.ingredients_label.grid(row=2, column=0)
        self.show_ingredients_text = tk.Text(self.line1_frame, wrap=WORD, width=50)
        self.show_ingredients_text.grid(row=2, column=1)

        # forth row
        self.show_instructions = tk.StringVar()
        self.show_instructions_label = tk.Label(self.line1_frame, textvariable=self.show_instructions)
        self.show_instructions.set("Instructions: ")
        self.show_instructions_label.grid(row=3, column=0)
        self.show_instructions_text = tk.Text(self.line1_frame, wrap=WORD, width=50)
        self.show_instructions_text.grid(row=3, column=1)
        self.addPrintButton = tk.Button(self.line1_frame, text='Print Recipe', fg='green', command=self.printRecipe)
        self.addPrintButton.grid(row=3, column=2)

        # fifth - 7th rows for database login info
        self.db_uname_label = tk.Label(self.line1_frame, text="Database Username:")
        self.db_uname_label.grid(row=4, column=0)
        self.db_uname = tk.Entry(self.line1_frame, width=50)
        self.db_uname.grid(row=4, column=1)
        self.db_password_label = tk.Label(self.line1_frame, text="Database Password:")
        self.db_password_label.grid(row=5, column=0)
        self.db_password = tk.Entry(self.line1_frame, width=50)
        self.db_password.grid(row=5, column=1)
        self.db_name_label = tk.Label(self.line1_frame, text="Database Name:")
        self.db_name_label.grid(row=6, column=0)
        self.db_name = tk.Entry(self.line1_frame, width=50)
        self.db_name.grid(row=6, column=1)

        # Packing the frame in the window
        self.line1_frame.pack()
        tk.mainloop()

    def viewRecipe(self):
        recipe_name = self.recipe_combo_box.get()
        self.show_ingredients_text.delete("1.0", END)
        self.show_instructions_text.delete("1.0", END)

        cnx = mysql.connector.connect(user='owen', password='password', database='recipes')
        cursor = cnx.cursor()

        # fetches and displays ingredients
        query = "SELECT ingredients FROM recipes WHERE name = %s"
        cursor.execute(query, (recipe_name,))
        ingredient_list = cursor.fetchone()
        self.show_ingredients_text.insert(END, ingredient_list[0])

        # fetches and displays instructions
        query1 = "SELECT instructions FROM recipes WHERE name = %s"
        cursor.execute(query1, (recipe_name,))
        instructions = cursor.fetchone()
        self.show_instructions_text.insert(END, instructions[0])

    def viewImage(self):

        try:

            recipe_name = self.recipe_combo_box.get()
            cnx = mysql.connector.connect(user='owen', password='password', database='recipes')
            cursor = cnx.cursor()
            query = 'SELECT photo FROM recipes WHERE name = %s'
            cursor.execute(query, (recipe_name,))

            image_data = cursor.fetchall()

            image = image_data[0][0]
            binary_data = base64.b64decode(image)
            image1 = PIL.Image.open(io.BytesIO(binary_data))
            image_list.append(image1)
            image_displayed = image_list[0]

            # image = PIL.Image.open(binary_data)
            resized_image = image_displayed.resize((512, 288), Image.LANCZOS)
            new_image = ImageTk.PhotoImage(resized_image)
            image_list.append(new_image)
            img = Label(self.line1_frame, image=new_image)
            img.grid(row=2, column=2)

            cursor.close()
            cnx.close()
        except TypeError:
            message = "No picture for this recipe."
            Label(self.line1_frame, text=message).grid(row=2, column=2)

    def getRecipes(self):
        ingredient = self.search_by_ingredient_input.get()
        cnx = mysql.connector.connect(user='owen', password='password', database='recipes')
        cursor = cnx.cursor()
        recipe_list = []

        query = "SELECT name FROM recipes WHERE ingredients LIKE %s ORDER BY name ASC"
        cursor.execute(query, ('%' + ingredient + '%',))
        for (name) in cursor:
            recipe_list.append(name)

        print(recipe_list)
        self.recipe_combo_box['values'] = recipe_list

        cursor.close()
        cnx.close()

    def printRecipe(self):
        document = Document()
        document.add_heading(self.recipe_combo_box.get().title(), 0)
        #document.add_picture('', width=Inches(3))
        document.add_heading('Ingredients', level=1)
        document.add_paragraph(self.show_ingredients_text.get(1.0, END))
        document.add_heading('Instructions', level=1)
        document.add_paragraph(self.show_instructions_text.get(1.0, END))
        document.save(self.recipe_combo_box.get() + '.docx')
        # os.system('start ' + self.name_entry.get() + '.docx')

        filename = self.recipe_combo_box.get() + '.docx'
        win32api.ShellExecute(0, "print", filename, '/c:"%s"' % win32print.GetDefaultPrinter(), ".", 1)


ViewRecipesUI()
