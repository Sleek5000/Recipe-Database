import base64
import tkinter as tk

from tkinter import *
from Recipe import Recipe
import mysql.connector
from docx import Document
import os
import win32print
import win32api
from tkinter import filedialog
from docx.shared import Inches


class AddRecipesUI:
    def __init__(self):
        self.master = tk.Tk()
        self.master.title('Add Recipes')
        self.master.geometry('1920x1080')

        self.menu = tk.Menu(self.master)
        self.master.config(menu=self.menu)
        self.navigation_menu = tk.Menu(self.menu)
        self.menu.add_cascade(label='Navigation', menu=self.navigation_menu)
        self.navigation_menu.add_command(label='Add Recipes')
        self.navigation_menu.add_command(label='View Recipes')
        self.navigation_menu.add_separator()
        self.navigation_menu.add_command(label='Exit', command=self.master.quit)

        self.line1_frame = tk.Frame(self.master)
        self.line2_frame = tk.Frame(self.master)
        self.line3_frame = tk.Frame(self.master)
        self.line4_frame = tk.Frame(self.master)
        self.line5_frame = tk.Frame(self.master)
        self.line6_frame = tk.Frame(self.master)
        self.line7_frame = tk.Frame(self.master)
        self.line8_frame = tk.Frame(self.master)

        self.recipe_label = tk.Label(self.line1_frame, text='Recipe Name')  # frame 1
        self.name_entry = tk.Entry(self.line1_frame)
        self.recipe_label.pack(side='left')
        self.name_entry.pack(side='left')

        self.course_label = tk.Label(self.line2_frame, text='Course')  # frame 2
        self.course_entry = tk.Listbox(self.line2_frame, selectmode=SINGLE)
        self.course_entry.insert(1, 'Breakfast')
        self.course_entry.insert(2, 'Lunch')
        self.course_entry.insert(3, 'Dinner')
        self.course_entry.insert(4, 'Appetizer')
        self.course_entry.insert(5, 'Dessert')
        self.course_label.pack(side='left')
        self.course_entry.pack(side='left')

        self.ingredients_label = tk.Label(self.line3_frame, text='Ingredients')  # frame 3
        self.ingredients_entry = tk.Text(self.line3_frame, wrap=WORD)
        self.ingredients_label.pack(side='left')
        self.ingredients_entry.pack(side='left')

        self.instruction_label = tk.Label(self.line4_frame, text='Instructions')  # frame 4
        self.instruction_entry = tk.Text(self.line4_frame, wrap=WORD)
        self.instruction_label.pack(side='left')
        self.instruction_entry.pack(side='left')

        self.time_label = tk.Label(self.line5_frame, text='Time to make')  # frame 5
        self.time_entry = tk.Entry(self.line5_frame)
        self.time_label.pack(side='left')
        self.time_entry.pack(side='left')

        self.portions_label = tk.Label(self.line6_frame, text='Portions')  # frame 6
        self.portions_entry = tk.Spinbox(self.line6_frame, from_=1, to=50)
        self.portions_label.pack(side='left')
        self.portions_entry.pack(side='left')

        # frame 7
        self.picture_to_be_added = ""
        self.label_file_explorer = Label(self.line7_frame,
                                    text="Recipe picture to upload.",
                                    fg="blue")
        self.button_explore = Button(self.line7_frame, text="Browse Files", command=self.browseFiles)
        self.label_file_explorer.pack(side='left')
        self.button_explore.pack(side='left')

    # self.nam_entry = tk.Entry(self.master)
        # self.course_entry = tk.Listbox(self.master, selectmode=SINGLE)
        # self.course_entry.insert(1, 'Breakfast')
        # self.course_entry.insert(2, 'Lunch')
        # self.course_entry.insert(3, 'Dinner')
        # self.course_entry.insert(4, 'Appetizer')
        # self.course_entry.insert(5, 'Dessert')
        # self.ingredients_entry = tk.Text(self.master)
        # self.instruction_entry = tk.Text(self.master)
        # self.time_entry = tk.Entry(self.master)
        # self.portions_entry = tk.Spinbox(self.master, from_=1, to=50)
        # self.name_entry.grid(row=0, column=1)
        # self.course_entry.grid(row=1, column=1)
        # self.ingredients_entry.grid(row=2, column=1)
        # self.instruction_entry.grid(row=3, column=1)
        # self.time_entry.grid(row=4, column=1)
        # self.portions_entry.grid(row=5, column=1)
        self.addRecipeButton = tk.Button(self.line8_frame, text='Add Recipe', fg='blue', command=self.addRecipe)
        self.addRecipeButton.pack(side='left')
        self.addPrintButton = tk.Button(self.line8_frame, text='Print Recipe', fg='green', command=self.printRecipe)
        self.addPrintButton.pack(side='right')

        self.line1_frame.pack()
        self.line2_frame.pack()
        self.line3_frame.pack()
        self.line4_frame.pack()
        self.line5_frame.pack()
        self.line6_frame.pack()
        self.line7_frame.pack()
        self.line8_frame.pack()

        tk.mainloop()

    def printRecipe(self):
        document = Document()
        document.add_heading(self.name_entry.get().title(), 0)
        document.add_picture(self.picture_to_be_added, width=Inches(3))
        document.add_heading('Ingredients', level=1)
        document.add_paragraph(self.ingredients_entry.get(1.0, END))
        document.add_heading('Instructions', level=1)
        document.add_paragraph(self.instruction_entry.get(1.0, END))
        document.save(self.name_entry.get()+'.docx')
        #os.system('start ' + self.name_entry.get() + '.docx')

        filename = self.name_entry.get()+'.docx'
        win32api.ShellExecute(0, "print", filename, '/c:"%s"' % win32print.GetDefaultPrinter(), ".", 1)

    def addRecipe(self):
        name = self.name_entry.get()
        course = str(self.course_entry.get(ACTIVE))
        ingredients = self.ingredients_entry.get(1.0, END)
        instructions = self.instruction_entry.get(1.0, END)
        time = self.time_entry.get()
        portions = int(self.portions_entry.get())

        recipe = Recipe(name, course, ingredients, instructions, time, portions)

        cnx = mysql.connector.connect(user='owen', password='password', database='recipes')
        cursor = cnx.cursor()

        query = """INSERT INTO recipes (name, course, ingredients, instructions, time, portions, photo)
             VALUES (%s, %s, %s, %s, %s, %s, %s)"""

        cursor.execute(query, (recipe.name,
                               recipe.course,
                               recipe.ingredients,
                               recipe.instructions,
                               recipe.time,
                               recipe.portions,
                               self.convertToBinaryData()))
        cnx.commit()

        cursor.close()
        cnx.close()

    def convertToBinaryData(self):
        with open(self.picture_to_be_added, 'rb') as file:
            binary_data = file.read()
            encodestring = base64.b64encode(binary_data)
            return encodestring

    def browseFiles(self):
        filename = filedialog.askopenfilename(initialdir="/",
                                              title="Select a file",
                                              filetypes=[("JPG", "*.jpg"),
                                                         ("JPEG", "*.jpeg"),
                                                         ("PNG", "*.PNG"),
                                                         ("ICON", "*.ico")])
        self.label_file_explorer.configure(text="Picture to be uploaded: " + filename)
        self.picture_to_be_added = filename


AddRecipesUI()
